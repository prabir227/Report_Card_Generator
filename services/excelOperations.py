from docx import Document
from openpyxl import load_workbook
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
class ExcelOperations:
    def __init__(self,lvl0Start, lvl0End, lvl1Start, lvl1End, lvl2Start, lvl2End, lvl3Start, lvl3End, lvl4Start, lvl4End, lvl5Start, lvl5End, date="15/06/28",eng1="C", eng2="K", eng3="S", hin1="D", hin2="L", hin3="T", math1="E", math2="M", math3="U", sst1="G", sst2="O", sst3="W", evs1="F", evs2="N", evs3="V", name="A",totalLvl0=150, totalLvl1=150, totalLvl2=150, totalLvl3=150, totalLvl4=150, totalLvl5=200 , centre="Boring Road"):
        self.__eng1 = eng1
        self.__eng2 = eng2
        self.__eng3 = eng3
        self.__hin1 = hin1
        self.__hin2 = hin2
        self.__hin3 = hin3
        self.__math1 = math1
        self.__math2 = math2
        self.__math3 = math3
        self.__sst1 = sst1
        self.__sst2 = sst2
        self.__sst3 = sst3
        self.__evs1 = evs1
        self.__evs2 = evs2
        self.__evs3 = evs3
        self.__totalLvl0 = totalLvl0
        self.__totalLvl1 = totalLvl1
        self.__totalLvl2 = totalLvl2
        self.__totalLvl3 = totalLvl3
        self.__totalLvl4 = totalLvl4
        self.__totalLvl5 = totalLvl5
        self.__centre = centre
        self.__lvl0Start = lvl0Start
        self.__lvl0End = lvl0End
        self.__lvl1Start = lvl1Start
        self.__lvl1End = lvl1End
        self.__lvl2Start = lvl2Start
        self.__lvl2End = lvl2End
        self.__lvl3Start = lvl3Start
        self.__lvl3End = lvl3End
        self.__lvl4Start = lvl4Start
        self.__lvl4End = lvl4End
        self.__lvl5Start = lvl5Start
        self.__lvl5End = lvl5End
        self.__date = date
        self.__name = name
        self.__doc = Document("assets/resultFormat.docx")
        self.__result = load_workbook("uploads/uploaded.xlsx")
        self.__marksTable = self.__doc.tables[1]
        self.__nameTable = self.__doc.tables[0]

        self.docxWriter()

    def write_study_centre(self, studyCentre):
        para = self.__doc.paragraphs[2]
        for run in para.runs:
            if "Study Centre: Vijay Nagar" in run.text:
                run.text = run.text.replace("Study Centre: Vijay Nagar", f"Study Centre: {studyCentre}")
    def is_number(self,value):
        try:
            float(value)
            return True
        except (ValueError, TypeError):
            return False
    def write_name_and_lvl_to_cell(self,table, name, lvl):
        #Writing name
        cell = table.cell(0, 0)
        cell.width = Inches(5)
        cell._tc.width = Inches(5)
        cell.text = ""
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run(f"Name  : {name}")
        run.font.size = Pt(16)
        run.bold = True

        #Writing level
        cell2 = table.cell(0, 2)
        cell2.width = Inches(1.2)
        cell2._tc.width = Inches(1.2)
        cell2.text = ""
        paragraph2 = cell2.paragraphs[0]
        run2 = paragraph2.add_run(f"Level : {lvl}")
        run2.font.size = Pt(16)
        run2.bold = True
    def write_to_cell(self,table, row_idx, col_idx, data):
        cell = table.cell(row_idx, col_idx)
        cell.text = ""  # Clear any existing text

        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = paragraph.add_run(str(data))
        run.font.size = Pt(14)
        run.bold = True
    def write_date(self, date):
        para = self.__doc.paragraphs[12]
        for run in para.runs:
            if run.text.strip() == "12/05/2025":
                run.text = date
                break
    def sstWriter(self,sheet,row,total):
        sst1=sheet[f"{self.__sst1}{row}"].value
        sst2=sheet[f"{self.__sst2}{row}"].value
        sst3=sheet[f"{self.__sst3}{row}"].value
        sstsum = 0.0
        # SST
        if self.is_number(sst1):
            sst1 = float(sst1)
            self.write_to_cell(self.__marksTable, 5, 1, sst1)
            sstsum += sst1
        else:
            self.write_to_cell(self.__marksTable, 5, 1, "N/A")

        if self.is_number(sst2):
            sst2 = float(sst2)
            self.write_to_cell(self.__marksTable, 5, 2, sst2)
            sstsum += sst2
        else:
            self.write_to_cell(self.__marksTable, 5, 2, "N/A")

        if self.is_number(sst3):
            sst3 = float(sst3)
            self.write_to_cell(self.__marksTable, 5, 3, sst3)
            sstsum += sst3
        else:
            self.write_to_cell(self.__marksTable, 5, 3, "N/A")

        self.write_to_cell(self.__marksTable, 5, 5, sstsum)
        self.write_to_cell(self.__marksTable, 5, 4, total)
        self.write_to_cell(self.__marksTable, 5, 6, f"{(sstsum / total) * 100:.2f}%")
        
    def englishWriter(self,sheet,row,total):
        eng1=sheet[f"{self.__eng1}{row}"].value
        eng2=sheet[f"{self.__eng2}{row}"].value
        eng3=sheet[f"{self.__eng3}{row}"].value
        engsum = 0.0
        # English
        if self.is_number(eng1):
            eng1 = float(eng1)
            self.write_to_cell(self.__marksTable, 1, 1, eng1)
            engsum += eng1
        else:
            self.write_to_cell(self.__marksTable, 1, 1, "N/A")

        if self.is_number(eng2):
            eng2 = float(eng2)
            self.write_to_cell(self.__marksTable, 1, 2, eng2)
            engsum += eng2
        else:
            self.write_to_cell(self.__marksTable, 1, 2, "N/A")

        if self.is_number(eng3):
            eng3 = float(eng3)
            self.write_to_cell(self.__marksTable, 1, 3, eng3)
            engsum += eng3
        else:
            self.write_to_cell(self.__marksTable, 1, 3, "N/A")
        
        self.write_to_cell(self.__marksTable, 1, 5, engsum)
        self.write_to_cell(self.__marksTable, 1, 4, total)
        self.write_to_cell(self.__marksTable, 1, 6, f"{(engsum / total) * 100:.2f}%")
    
    def hindiWriter(self, sheet, row, total):
        hin1 = sheet[f"{self.__hin1}{row}"].value
        hin2 = sheet[f"{self.__hin2}{row}"].value
        hin3 = sheet[f"{self.__hin3}{row}"].value
        hinsum = 0.0
        # Hindi
        if self.is_number(hin1):
            hin1 = float(hin1)
            self.write_to_cell(self.__marksTable, 2, 1, hin1)
            hinsum += hin1
        else:
            self.write_to_cell(self.__marksTable, 2, 1, "N/A")

        if self.is_number(hin2):
            hin2 = float(hin2)
            self.write_to_cell(self.__marksTable, 2, 2, hin2)
            hinsum += hin2
        else:
            self.write_to_cell(self.__marksTable, 2, 2, "N/A")

        if self.is_number(hin3):
            hin3 = float(hin3)
            self.write_to_cell(self.__marksTable, 2, 3, hin3)
            hinsum += hin3
        else:
            self.write_to_cell(self.__marksTable, 2, 3, "N/A")
        
        self.write_to_cell(self.__marksTable, 2, 5, hinsum)
        self.write_to_cell(self.__marksTable, 2, 4, total)
        self.write_to_cell(self.__marksTable, 2, 6, f"{(hinsum / total) * 100:.2f}%")

    def mathWriter(self,sheet,row,total):
        math1 = sheet[f"{self.__math1}{row}"].value
        math2 = sheet[f"{self.__math2}{row}"].value
        math3 = sheet[f"{self.__math3}{row}"].value
        mathsum = 0.0

        # Math
        if self.is_number(math1):
            math1 = float(math1)
            self.write_to_cell(self.__marksTable, 3, 1, math1)
            mathsum += math1
        else:
            self.write_to_cell(self.__marksTable, 3, 1, "N/A")

        if self.is_number(math2):
            math2 = float(math2)
            self.write_to_cell(self.__marksTable, 3, 2, math2)
            mathsum += math2
        else:
            self.write_to_cell(self.__marksTable, 3, 2, "N/A")

        if self.is_number(math3):
            math3 = float(math3)
            self.write_to_cell(self.__marksTable, 3, 3, math3)
            mathsum += math3
        else:
            self.write_to_cell(self.__marksTable, 3, 3, "N/A")

        self.write_to_cell(self.__marksTable, 3, 5, mathsum)
        self.write_to_cell(self.__marksTable, 3, 4, total)
        self.write_to_cell(self.__marksTable, 3, 6, f"{(mathsum / total) * 100:.2f}%")

    def evsWriter(self,sheet, row, total):
        evs1 = sheet[f"{self.__evs1}{row}"].value
        evs2 = sheet[f"{self.__evs2}{row}"].value
        evs3 = sheet[f"{self.__evs3}{row}"].value

        evssum = 0.0
        # EVS
        if self.is_number(evs1):
            evs1 = float(evs1)
            self.write_to_cell(self.__marksTable, 4, 1, evs1)
            evssum += evs1
        else:
            self.write_to_cell(self.__marksTable, 4, 1, "N/A")

        if self.is_number(evs2):
            evs2 = float(evs2)
            self.write_to_cell(self.__marksTable, 4, 2, evs2)
            evssum += evs2
        else:
            self.write_to_cell(self.__marksTable, 4, 2, "N/A")

        if self.is_number(evs3):
            evs3 = float(evs3)
            self.write_to_cell(self.__marksTable, 4, 3, evs3)
            evssum += evs3
        else:
            self.write_to_cell(self.__marksTable, 4, 3, "N/A")


        # Totals & Percentages
        self.write_to_cell(self.__marksTable, 4, 5, evssum)
        self.write_to_cell(self.__marksTable, 4, 4, total)
        self.write_to_cell(self.__marksTable, 4, 6, f"{(evssum / total) * 100:.2f}%")

    def docxWriterLoop(self, startRow, endRow, lvl, total):
        sheet = self.__result.active
        for row in range(startRow, endRow + 1):
            name = sheet[f"{self.__name}{row}"].value
            evssum = 0.0
            self.write_name_and_lvl_to_cell(self.__nameTable, name, lvl)
            self.write_study_centre(self.__centre)
            self.write_date(self.__date)

            #English
            self.englishWriter(sheet, row, total)
            # Hindi
            self.hindiWriter(sheet, row, total)

            # Math
            self.mathWriter(sheet, row, total)

            # EVS
            self.evsWriter(sheet, row, total)

            if(lvl==5):
                self.sstWriter(sheet,row,total)

            self.__doc.save(f"output/lvl{lvl}/{name}.docx")

    def docxWriter(self):
        
        self.docxWriterLoop(self.__lvl0Start, self.__lvl0End, 0, self.__totalLvl0)
        self.docxWriterLoop(self.__lvl1Start, self.__lvl1End, 1, self.__totalLvl1)
        self.docxWriterLoop(self.__lvl2Start, self.__lvl2End, 2, self.__totalLvl2)
        self.docxWriterLoop(self.__lvl3Start, self.__lvl3End, 3, self.__totalLvl3)
        self.docxWriterLoop(self.__lvl4Start, self.__lvl4End, 4, self.__totalLvl4)
        self.docxWriterLoop(self.__lvl5Start, self.__lvl5End, 5, self.__totalLvl5)
    



        

